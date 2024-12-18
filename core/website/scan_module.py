import os
import json
import pandas as pd
import io
import csv
import re
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from pathlib import Path
import os
import json
import pandas as pd
import json
import io
import csv
import re
from openpyxl import load_workbook
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.shared import Pt, Cm
from pathlib import Path
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER

def load_keywords_from_json():
    """
    Load keywords from a JSON file relative to the current script.

    Returns:
    - dict: The loaded keyword dictionary from the JSON file.
    """
    current_directory = os.path.dirname(os.path.abspath(__file__))
    json_file = os.path.join(current_directory, 'keywords.json')

    try:
        with open(json_file, 'r', encoding='utf-8') as file:
            data = json.load(file)
    except FileNotFoundError:
        raise FileNotFoundError(f"File not found: {json_file}")

    return data

def extract_and_iterate_docx_content(file_path, table_id=None, **pandas_kwargs):
    """
    Extracts all content from a DOCX file, including headers, footers, paragraphs, and tables.

    Parameters:
    - file_path: Path to the DOCX file.
    - table_id: (Optional) Index of the specific table to extract. If None, all tables will be extracted.
    - pandas_kwargs: Optional keyword arguments passed to `pd.read_csv()` for parsing tables.

    Returns:
    - list: List of extracted content (paragraphs, tables, headers) as lowercase strings.
    """
    document = Document(file_path)
    content_list = []

    def extract_single_table(table):
        memory_file = io.StringIO()
        csv_writer = csv.writer(memory_file)

        for row in table.rows:
            csv_writer.writerow([cell.text.strip() for cell in row.cells])

        memory_file.seek(0)
        return pd.read_csv(memory_file, **pandas_kwargs)

    for child_element in document.element.body.iterchildren():
        if isinstance(child_element, CT_P):
            paragraph = Paragraph(child_element, document).text.strip().lower()
            if paragraph:
                content_list.append(paragraph)
        elif isinstance(child_element, CT_Tbl):
            table = Table(child_element, document)
            if table_id is None or document.tables.index(table) == table_id:
                table_df = extract_single_table(table)
                content_list.append(table_df.to_string(index=False).lower())

    for section in document.sections:
        for header in section.header.paragraphs:
            header_text = header.text.strip().lower()
            if header_text:
                content_list.append(header_text)

    return content_list

def find_keywords_and_patterns_in_docx(content_list, keywords_dict, patterns):
    """
    Search for both keywords and regex patterns in DOCX content, and return the count of occurrences.

    Parameters:
    - content_list: List of content extracted from the DOCX file.
    - keywords_dict: Dictionary of keywords to search for.
    - patterns: Dictionary of regex patterns to search for.

    Returns:
    - str: JSON formatted string containing the search results for keywords and patterns.
    """
    found_keywords = []
    found_patterns = []
    keyword_counts = {}
    pattern_counts = {}

    docx_content_lower = " ".join([str(content).strip().lower() for content in content_list])

    for keyword_list in keywords_dict.values():
        for keyword in keyword_list:
            keyword_lower = keyword.lower().strip()
            keyword_counts[keyword_lower] = docx_content_lower.count(keyword_lower)
            if keyword_counts[keyword_lower] > 0:
                found_keywords.append({
                    "Found Keyword": keyword,
                    "num of the same keyword": keyword_counts[keyword_lower]
                })

    for pattern_name, pattern_regex in patterns.items():
        matches = pattern_regex.findall(docx_content_lower)
        pattern_counts[pattern_name] = len(matches)
        if pattern_counts[pattern_name] > 0:
            found_patterns.append({
                "Pattern Name": pattern_name,
                "num of the same pattern": pattern_counts[pattern_name]
            })

    result = {
        "Keywords": found_keywords,
        "Patterns": found_patterns
    }

    return json.dumps(result, indent=4, ensure_ascii=False)

def check_keywords_and_patterns_in_docx(file_path, patterns):
    """
    Check the DOCX content for both keywords and regex patterns.

    Parameters:
    - file_path: Path to the DOCX file.
    - patterns: Dictionary of regex patterns to search for.

    Returns:
 - str: JSON formatted string containing the search results.
    """
    keywords_dict = load_keywords_from_json()
    content_list = extract_and_iterate_docx_content(file_path)
    result = find_keywords_and_patterns_in_docx(content_list, keywords_dict, patterns)
    return result

def extract_and_iterate_excel_content(file_path, **pandas_kwargs):
    """
    Extract content from all sheets of an Excel file, including column headers.

    Parameters:
    - file_path: Path to the Excel file.
    - pandas_kwargs: Optional keyword arguments passed to `pd.read_excel()`.

    Returns:
    - dict: Dictionary where each sheet name maps to a list of extracted content.
    """
    sheet_dict = pd.read_excel(file_path, sheet_name=None, dtype=str, **pandas_kwargs)
    content_dict = {}

    for sheet_name, df in sheet_dict.items():
        content_list = []
        headers = df.columns
        content_list.extend([str(header).strip() for header in headers if "Unnamed" not in header])

        for _, row in df.iterrows():
            for cell in row:
                cell_str = str(cell).strip()
                if pd.notnull(cell) and "Unnamed" not in cell_str:
                    content_list.append(cell_str)

        content_dict[sheet_name] = content_list

    return content_dict

def find_keywords_and_patterns_in_excel(excel_content_dict, keywords_dict, patterns):
    """
    Search for both keywords and regex patterns in Excel content.

    Parameters:
    - excel_content_dict: Dictionary of content from Excel sheets.
    - keywords_dict: Dictionary of keywords to search for.
    - patterns: Dictionary of regex patterns to search for.

    Returns:
    - str: JSON formatted string containing search results.
    """
    found_keywords = []
    found_patterns = []

    for sheet_name, excel_content in excel_content_dict.items():
        keyword_counts = {}
        pattern_counts = {}

        for keyword_list in keywords_dict.values():
            for keyword in keyword_list:
                keyword_lower = keyword.lower().strip()
                keyword_counts[keyword_lower] = sum(content.lower().count(keyword_lower) for content in excel_content)

                if keyword_counts[keyword_lower] > 0:
                    found_keywords.append({
                        "Sheet": sheet_name,
                        "Found Keyword": keyword,
                        "num of the same keyword": keyword_counts[keyword_lower]
                    })

        for pattern_name, pattern_regex in patterns.items():
            pattern_counts[pattern_name] = sum(len(pattern_regex.findall(content)) for content in excel_content)

            if pattern_counts[pattern_name] > 0:
                found_patterns.append({
                    "Sheet": sheet_name,
                    "Pattern Name": pattern_name,
                    "num of the same pattern": pattern_counts[pattern_name]
                })

    result = {
        "Keywords": found_keywords,
        "Patterns": found_patterns
    }

    return json.dumps(result, indent=4, ensure_ascii=False)

def check_keywords_and_patterns_in_excel(xlsx_file, patterns):
    """
    Check an Excel file for both keywords and regex patterns.

    Parameters:
    - xlsx_file: Path to the Excel file.
    - patterns: Dictionary of regex patterns to search for.

    Returns:
    - str: JSON formatted string containing search results.
    """
    keywords_dict = load_keywords_from_json()
    excel_content_dict = extract_and_iterate_excel_content(xlsx_file)
    result = find_keywords_and_patterns_in_excel(excel_content_dict, keywords_dict, patterns)
    return result

def define_rules():
    """
    Define a set of rules consisting of various keywords to search for.

    Returns:
    - dict: Dictionary of rule sets where each rule contains relevant keywords.
    """
    return {
        "rule_1": {
            "email": "email",
            "id number (cccd/cmnd)": "id number (cccd/cmnd)",
            "tên khách hàng": [
                "tên khách hàng", "họ tên", "họ và tên", "họ tên", "name", "full name"
            ],
            "địa chỉ": "địa chỉ",
            "số điện thoại": "số điện thoại"
        },
        "rule_2": {
            "số thẻ": "số thẻ",
            "cvv": "cvv",
            "ngày hết hạn": "ngày hết hạn",
            "tên chủ thẻ": [
                "tên khách hàng", "chủ thẻ", "tên chủ thẻ", "họ và tên", "họ tên", "name", "full name"
            ]
        },
        "rule_3": {
            "Thông tin về chủ trương đầu tư Dự án CNTT": {
                "key1": [
                    "dự án", "mục tiêu đầu tư", "sự cần thiết", "phương án đầu tư", 
                    "hạng mục/cấu phần mua sắm", "tuân thủ kiến trúc", "phương án kỹ thuật sơ bộ", 
                    "khái toán", "hiệu quả đầu tư", "báo giá"
                ]
            },
            "Thông tin về tiêu chuẩn kinh tế kỹ thuật của Dự án CNTT": {
                "key2": [
                    "kinh tế kỹ thuật", "dự án", "căn cứ pháp lý", "nội dung dự án", 
                    "mục tiêu đầu tư", "tổng mức đầu tư", "mức độ tuân thủ kiến trúc", 
                    "yêu cầu về làm chủ", "hình thức mua sắm bản quyền", "cấp độ hệ thống thông tin", 
                    "mức độ kiểm soát", "rủi ro", "tiêu chuẩn kỹ thuật", "dự toán chi tiết", 
                    "kế hoạch lựa chọn nhà thầu", "giá gói thầu", "phương thức lựa chọn nhà thầu"
                ]
            },
            "Thông tin mời thầu của Dự án CNTT": {
                "key3": [
                    "hồ sơ mời thầu", "số hiệu gói thầu", "tên gói thầu", "thủ tục đấu thầu", 
                    "yêu cầu về kỹ thuật", "biểu mẫu hợp đồng", "chỉ dẫn nhà thầu", 
                    "bảng dữ liệu đầu thầu", "biểu mẫu mời thầu và dự thầu", 
                    "tiêu chuẩn đánh giá về kỹ thuật", "năng lực và kinh nghiệm"
                ]
            }
        },
        "rule_4": {
            "or_1": ["dntd bq", "dntd ck"],
            "or_2": ["số lượng sản phẩm", "slsp"],
            "or_3": ["số lượng khách hàng", "số lượng kh", "slkh"],
            "or_4": ["hđv bq", "hđv ck"],
            "and": ["tnt", "nim td", "nim hđv", "cir", "cltc"]
        },
        "rule_5": [
            "etl", "tài liệu mapping chi tiết"
        ]
    }

def scan_file(file_path):
    """
    Scan a file and check for both keywords and patterns depending on the file type.

    Parameters:
    - file_path: Path to the file to be scanned.

    Returns:
    - dict: JSON object with the scan results.
    - str: Error or success message.
    """
    if not file_path:
        return None, "Vui lòng chọn một file trước khi scan."

    file_extension = Path(file_path).suffix.lower()
    patterns = {
        "stk 10 số chứa 3 BDS": re.compile(r'\b(111|116|117|118|119|120|121|122|123|124|125|126|128|129|130|131|132|133|134|135|136|138|139|140|141|144|145|147|149|150|151|159|160|166|168|169|177|180|181|186|188|189|199|211|212|213|214|215|216|217|220|222|256|260|261|268|279|289|310|311|313|314|315|317|318|319|321|328|330|341|345|351|362|368|371|375|376|390|395|398|411|421|425|426|427|428|431|432|433|440|441|443|444|448|450|451|452|455|460|461|465|466|468|471|480|482|483|486|488|501|502|505|512|513|518|520|522|531|532|540|556|558|560|561|562|566|565|570|573|580|581|590|601|602|611|615|620|621|625|631|632|633|635|636|641|642|646|650|651|652|653|655|656|661|670|671|672|679|680|686|691|696|701|702|710|711|721|729|730|735|737|741|742|745|748|750|753|760|761|762|766|780|785|788)\d{7}\b'),
        "cvv": re.compile(r"(?i)(\bCVV\b[\s\S]*?\b\d{3}\b)"),
        "cvv": re.compile(r"(?i)(\bCVV\b[\s\S]*?\b\d{3}\b)"),
        "id number (cccd/cmnd)": re.compile(r'\b(?:cmnd|cccd|CMND|CCCD)\b.*\b\d{9}\b|\b(?:cmnd|cccd|CMND|CCCD)\b.*\b\d{12}\b|\b\d{9}\b|\b\d{12}\b'),
        "địa chỉ": re.compile(r'(\d+)\s+(đường|phố|phường|quận|thành phố)\s+(\w+),\s+(\w+),\s+(\w+)'),
        "giá trị tiền tối thiểu 6 chữ số": re.compile(r'\b\d{1,3}(,\d{3}){1,}\b'),
        "số điện thoại": re.compile(r'\b(0\d{9}|\+[\d]{11})\b'),
        "email": re.compile(r'\S+@\S+'),
        "số thẻ": re.compile(r'''
            \b                                        # Start of word boundary
            (?:9704|476632|411153|428695|427126|402460|406220|511957|
            517107|517453|542726|530515|515110|51511)  # Card prefixes
            (                                         # Matching options
                \d{10,15}                             # 10-15 digit numbers
                |\d{16}                               # 16 digit numbers
                |\d{19}                               # 19 digit numbers
                |\d{20}                               # 20 digit numbers
                |\d{3} \d{4} \d{4} \d{4} \d{4}        # 16 digits in 3-4-4-4 groups
                |\d{4} \d{4} \d{4} \d{4} \d{3}        # 19 digits in 4-4-4-4-3 groups
                |\d{4} \d{4} \d{4} \d{4} \d{4}        # 20 digits in 4-4-4-4-4 groups
            )
            \b                                        # End of word boundary
        ''', re.VERBOSE),
        "số tài khoản ẩn": re.compile(r'\b(?:9704|476632|411153|428695|427126|402460|406220|511957|517107|517453|542726|530515|515110)(?:\d{2}xxxx\d{4}|\d{3}xxxx\d{4}|\d{4}xxxx\d{4})\b')
    }

    results = None

    if file_extension == '.docx':
        results = check_keywords_and_patterns_in_docx(file_path, patterns)
    elif file_extension == '.xlsx':
        results = check_keywords_and_patterns_in_excel(file_path, patterns)
    elif file_extension in [".pdf", ".ptpx", "tsv", "txt", "py", "png", "jpg"]:
        results = "chưa làm"  # Not yet implemented for these types

    if results is None:
        return None, "Không có kết quả hợp lệ từ file đã chọn."

    results_json = json.loads(results)
    return results_json

def classify_document_with_multiple_rules(results, rules):
    """
    Classify a document based on multiple rules by checking keywords and patterns.

    Parameters:
    - results: Dictionary containing scan results with found keywords and patterns.
    - rules: Dictionary containing multiple rule sets to compare with results.

    Returns:
    - str: Classification label (e.g., 'Confidential', 'Public', etc.).
    - dict: A dictionary containing all matching rules and their details.
    """
    if results == "chưa làm" or results is None:
        return "chưa làm", {"error": "Hiện tại không hỗ trợ định dạng tệp."}

    if not isinstance(results, dict):
        return "Unsupported file", {"error": "Kết quả không hợp lệ."}

    if 'Keywords' not in results or not isinstance(results['Keywords'], list):
        return "Public", {"error": "Không tìm thấy từ khóa hợp lệ trong kết quả."}

    try:
        found_keywords = {item['Found Keyword'] for item in results['Keywords'] 
                          if isinstance(item, dict) and 'Found Keyword' in item}
        found_patterns = {item['Pattern Name']: item['num of the same pattern']
                          for item in results['Patterns'] 
                          if isinstance(item, dict) and 'Pattern Name' in item and 'num of the same pattern' in item}
    except TypeError as e:
        return "Internal", {"error": f"Đã xảy ra lỗi khi phân tích kết quả: {str(e)}"}

    def matches_keywords_rule(result_keywords, rule_keywords):
        if isinstance(rule_keywords, list):
            return any(keyword in result_keywords for keyword in rule_keywords)
        return rule_keywords in result_keywords

    results_dict = {}

    def process_rule_1():
        if 'rule_1' not in rules:
            return None
        
        rule_1_matches = 0
        matched_rule_1_keys = []
        matched_pattern_counts = []

        for key, value in rules['rule_1'].items():
            keyword_match = matches_keywords_rule(found_keywords, value)
            pattern_match = key in found_patterns and found_patterns[key] >= 10

            if keyword_match or pattern_match:
                rule_1_matches += 1
                matched_rule_1_keys.append(key)
                if pattern_match:
                    matched_pattern_counts.append(f"{key}: {found_patterns[key]}")

        return {
            "matched_keys": matched_rule_1_keys,
            "pattern_counts": matched_pattern_counts
        } if rule_1_matches >= 3 else None

    def process_rule_2():
        if 'rule_2' not in rules:
            return None
        
        matched_rule_2_keys = []

        for key, value in rules['rule_2'].items():
            keyword_match = matches_keywords_rule(found_keywords, value)
            pattern_match = key in found_patterns

            if not (keyword_match or pattern_match):
                return None
            matched_rule_2_keys.append(key)

        return {"matched_keys": matched_rule_2_keys}

    def process_rule_3():
        """
        Process Rule 3: Check for all keywords in subgroups.

        Parameters:
        - rules: Dictionary containing the rules.
        - found_keywords: Set of keywords found in the document.

        Returns:
        - dict: A dictionary containing matched keys and their details, or None if no matches found.
        """
        if 'rule_3' not in rules:
            return None
        
        matched_rule_3 = []

        # Iterate through each subgroup in rule_3
        for subgroup, subgroup_values in rules['rule_3'].items():
            # Initialize a flag to check if all keywords in the subgroup match
            all_keywords_matched = True
            
            # Iterate through each key and its associated values
            for key, values in subgroup_values.items():
                # Check if all keywords in the values are found in the found_keywords
                if not all(keyword in found_keywords for keyword in values):
                    all_keywords_matched = False
                    break  # No need to check further if any keyword is not matched

            # If all keywords in the subgroup matched, add to the results
            if all_keywords_matched:
                matched_rule_3.append({
                    "subgroup": subgroup,
                    "matched_values": [keyword for key, values in subgroup_values.items() for keyword in values]  # Return all values since they all matched
                })

        return {"matched_keys": matched_rule_3} if matched_rule_3 else None
    

    def process_rule_4():
        if 'rule_4' not in rules:
            return None
        
        def matches_all_keywords(rule_keywords):
            return all(keyword in found_keywords for keyword in rule_keywords)

        def matches_any_keywords(rule_keywords):
            return any(keyword in found_keywords for keyword in rule_keywords)

        rule_4_conditions = rules['rule_4']
        and_condition_met = matches_all_keywords(rule_4_conditions.get('and', []))
        or_conditions = ['or_1', 'or_2', 'or_3', 'or_4']
        or_condition_met = all(
            matches_any_keywords(rule_4_conditions.get(or_key, [])) for or_key in or_conditions
        )

        return {
            "conditions_met": {
                "and": rule_4_conditions['and'],
                "or_keys": {
                    or_key: rule_4_conditions[or_key] for or_key in or_conditions
                }
            }
        } if and_condition_met and or_condition_met else None

    def process_rule_5():
        if 'rule_5' not in rules:
            return None
        
        def matches_all_keywords(rule_keywords):
            return all(keyword in found_keywords for keyword in rule_keywords)

        rule_5_values = rules['rule_5']
        return {"matched_values": rule_5_values} if matches_all_keywords(rule_5_values) else None

    rule_processors = [
        process_rule_1,
        process_rule_2,
        process_rule_3,
        process_rule_4,
        process_rule_5
    ]

    for idx, processor in enumerate(rule_processors, 1):
        rule_result = processor()
        if rule_result:
            results_dict[f"rule_{idx}"] = rule_result

    return ("Confidential" if results_dict else "Internal", 
            json.dumps(results_dict, indent=4, ensure_ascii=False) if results_dict else {"": "Không có quy tắc nào được áp dụng."})

def label_docx_file(file_path, classify):
    document = Document(file_path)
    scan_result = scan_file(file_path)
    rul = define_rules()
    classify, sms = classify_document_with_multiple_rules(scan_result, rul)

    if classify == "Confidential":
        label_text = "Tài liệu Mật của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"
    elif classify == "Internal":
        label_text = "Tài liệu Nội bộ của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"
    else:
        label_text = ""

    section = document.sections[0]
    section.top_margin = Cm(2)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)

    footer = section.footer

    for element in footer._element.xpath('.//w:p'):
        element.getparent().remove(element)

    paragraph = footer.add_paragraph()
    right_margin_position = section.page_width - section.right_margin - section.left_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(right_margin_position, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

    run_label = paragraph.add_run(label_text)
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(12)

    run_label.add_tab()

    run_page = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)
    run_page._r.append(fldChar3)

    run_page.font.name = 'Times New Roman'
    run_page.font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    document.save(file_path)
    return label_text

def label_xlsx_file_footer(file_path, classify):
    scan_result = scan_file(file_path)
    rul = define_rules()
    classify, sms = classify_document_with_multiple_rules(scan_result, rul)

    if classify == "Confidential":
        label_text = "Tài liệu Mật của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"
    elif classify == "Internal":
        label_text = "Tài liệu Nội bộ của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"
    else:
        label_text = ""

    workbook = load_workbook(file_path)

    for sheet in workbook.worksheets:
        new_footer = f"&L{label_text} &RTrang &P"
        sheet.oddFooter.center.text = new_footer
        sheet.oddFooter.center.size = 12
        sheet.oddFooter.center.font = "Times New Roman"

    workbook.save(file_path)
    return label_text

def edit_label_docx_file(file_path, label_text):
    document = Document(file_path)
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)

    footer = section.footer

    for element in footer._element.xpath('.//w:p'):
        element.getparent().remove(element)

    paragraph = footer.add_paragraph()
    right_margin_position = section.page_width - section.right_margin - section.left_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(right_margin_position, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

    run_label = paragraph.add_run(label_text)
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(12)

    run_label.add_tab()

    run_page = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)
    run_page._r.append(fldChar3)

    run_page.font.name = 'Times New Roman'
    run_page.font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    document.save(file_path)
    return label_text

def edit_label_xlsx_file(file_path, label_text):
    workbook = load_workbook(file_path)

    for sheet in workbook.worksheets:
        new_footer = f"&L{label_text} &R &P"
        sheet.oddFooter.center.text = new_footer
        sheet.oddFooter.center.size = 12
        sheet.oddFooter.center.font = "Times New Roman"

    workbook.save(file_path)
    return label_text
